from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED
import shutil
import tempfile
import textwrap
from io import BytesIO

from docx import Document
from PIL import Image, ImageDraw, ImageFont


SOURCE = Path(r"C:\Users\skk88\Downloads\OMNIGRAPH_Project_Documentation.docx")
OUTPUT = Path(r"C:\Users\skk88\Downloads\AI_Data_Analyst_Agent_Project_Documentation.docx")

PROJECT = "AI Data Analyst Agent"
STUDENT = "Sankuri kiran kumar"
ROLL = "23EG106b52"


REPLACEMENTS = [
    ("OMNIGRAPH Intelligence Platform", "AI Data Analyst Agent"),
    ("OMNIGRAPH", "AI Data Analyst Agent"),
    ("AI-Powered Graph RAG Enterprise Knowledge System", "AI-Powered Data Analytics and Natural Language BI Platform"),
    ("Transforming Documents into Connected Intelligence Graph", "Turning Raw Data into Conversational Business Intelligence"),
    ("V Harishwar", STUDENT),
    ("23EG107F34", ROLL),
    ("Graph-RAG", "RAG Analytics"),
    ("Graph RAG", "RAG Analytics"),
    ("graph-based reasoning", "data-driven reasoning"),
    ("graph intelligence", "analytics intelligence"),
    ("Knowledge Graph", "Analytics Context Store"),
    ("knowledge graph", "analytics context store"),
    ("Neural Knowledge Network", "AI Analytics Workspace"),
    ("Neo4j Browser", "ChromaDB / API docs"),
    ("Neo4j Aura (Cloud Free Tier)", "ChromaDB Cloud or local Chroma server"),
    ("Neo4j Desktop", "local ChromaDB service"),
    ("https://neo4j.com/download/", "https://www.trychroma.com/"),
    ("https://neo4j.com/cloud/platform/aura-graph-database/", "https://www.trychroma.com/"),
    ("NEO4J_PASSWORD", "CHROMA_API_KEY"),
    ("NEO4J_URI", "CHROMA_HOST"),
    ("NEO4J_USER", "CHROMA_TENANT"),
    ("Create a new project, start the database", "Start a local ChromaDB service or configure hosted Chroma"),
    ("Set password in CHROMA_API_KEY env var", "Set CHROMA_API_KEY only when using hosted Chroma"),
    ("Copy connection URI to CHROMA_HOST env var", "Copy host and database values to Chroma environment variables"),
    ("Neo4j", "ChromaDB"),
    ("Cypher", "vector search and dataframe queries"),
    ("React Flow", "Plotly.js"),
    ("Graph Explorer", "Dataset Analytics Suite"),
    ("Discovery Panel", "Admin and Insights Panel"),
    ("Discovery Engine", "Insight Discovery Engine"),
    ("FutureScope Simulation Engine", "Executive Insight Engine"),
    ("FutureScope Engine", "Executive Insight Engine"),
    ("FutureScope", "Executive Insight"),
    ("Future Simulation APIs", "Reporting and Cleaning APIs"),
    ("Future Simulation", "Executive Reporting"),
    ("Future Scope", "Future Scope"),
    ("Node.js + Express", "FastAPI"),
    ("Node.js", "Python"),
    ("Express.js", "FastAPI"),
    ("Express API Server", "FastAPI Server"),
    ("Multer", "FastAPI UploadFile"),
    ("pdf-parse / pdfjs-dist", "Pandas, OpenPyXL, PyMuPDF, and python-docx"),
    ("pdf-parse", "Pandas / PyMuPDF"),
    ("pdfjs-dist", "OpenPyXL / python-docx"),
    ("PDFBox", "Pandas / OpenPyXL"),
    ("Prisma ORM", "SQLAlchemy ORM"),
    ("Prisma", "SQLAlchemy"),
    ("PostgreSQL (Neon)", "SQLite / PostgreSQL-ready database"),
    ("PostgreSQL", "SQLite / PostgreSQL"),
    ("OpenAI GPT-4o-mini", "Google Gemini"),
    ("OpenAI", "Google Gemini"),
    ("Groq Llama-3.3", "Groq Llama 3.3"),
    ("React 18 + Vite", "Next.js 16 + React 19"),
    ("React 18", "React 19"),
    ("Vite", "Next.js"),
    ("React Router 6", "Next.js App Router"),
    ("Axios", "Fetch API"),
    ("Glassmorphism dark-mode aesthetic", "responsive analytics dashboard aesthetic"),
    ("PDF Upload & Text Extraction", "Dataset Upload & File Parsing"),
    ("PDF Upload", "CSV / Excel / Document Upload"),
    ("Upload PDF", "Upload Dataset"),
    ("PDF file", "CSV, Excel, or document file"),
    ("PDF files", "CSV, Excel, PDF, TXT, and DOCX files"),
    ("PDF document", "dataset file"),
    ("PDF", "CSV / Excel / Document"),
    ("Document & Extraction APIs", "Dataset & Analysis APIs"),
    ("Document APIs", "Dataset APIs"),
    ("Upload Document", "Upload Dataset"),
    ("Get Documents", "List Datasets"),
    ("Delete Document", "Delete Dataset"),
    ("document ingestion", "dataset ingestion"),
    ("Document Ingestion", "Dataset Ingestion"),
    ("uploaded documents", "uploaded datasets"),
    ("processed documents", "processed datasets"),
    ("document metadata", "dataset metadata"),
    ("Document metadata", "Dataset metadata"),
    ("document source", "dataset source"),
    ("Entity Extraction", "Schema and Insight Extraction"),
    ("Relationship Extraction", "Metric and Pattern Extraction"),
    ("Knowledge Graph Generation", "Analytics Context Generation"),
    ("Graph Visualization", "Chart Visualization"),
    ("Graph APIs", "Dataset APIs"),
    ("Graph Intelligence APIs", "Analytics Intelligence APIs"),
    ("Get Complete Graph", "Get Dataset Profile"),
    ("Get Graph Nodes", "Preview Dataset Rows"),
    ("Get Graph Relationships", "Get Chat History"),
    ("Get Relationships", "Get Chat History"),
    ("Get Graph Statistics", "Get Admin Statistics"),
    ("AI Chat", "Analysis Chat"),
    ("Hidden Relationship Discovery", "Sentiment and Pattern Discovery"),
    ("hidden relationship", "hidden pattern"),
    ("Hidden Relationship", "Hidden Pattern"),
    ("relationship-aware", "data-aware"),
    ("relationship suggestions", "insight suggestions"),
    ("relationships", "patterns"),
    ("Relationships", "Patterns"),
    ("relationship", "pattern"),
    ("Relationship", "Pattern"),
    ("entities", "columns and insights"),
    ("Entities", "Columns and Insights"),
    ("entity", "data insight"),
    ("Entity", "Data Insight"),
    ("nodes", "records and insights"),
    ("Nodes", "Records and Insights"),
    ("node", "record"),
    ("Node", "Record"),
    ("edges", "analytics links"),
    ("Edges", "Analytics Links"),
    ("edge", "analytics link"),
    ("Edge", "Analytics Link"),
    ("typed labels", "typed metadata"),
    ("typed relationship", "typed metric"),
    ("typed relationships", "typed metrics"),
    ("typed nodes", "typed dataset records"),
    ("label-based", "metadata-based"),
    ("schema-less at the graph level", "flexible for different dataset schemas"),
    ("multi-hop", "multi-step"),
    ("Multi-Hop", "Multi-Step"),
    ("traversal", "retrieval"),
    ("traverses", "analyzes"),
    ("traversing", "analyzing"),
    ("connected intelligence", "conversational analytics"),
    ("connected knowledge", "data-driven insight"),
    ("connected relationships", "data patterns"),
    ("organizational knowledge", "organizational data"),
    ("organizational memory", "analytics memory"),
    ("Digital Brain Memory System", "Long-Term Analytics Memory"),
    ("Autonomous Enterprise Digital Twin", "Autonomous Analytics Workspace"),
    ("Autonomous Graph Agent", "Autonomous Data Analysis Agent"),
    ("Predictive Relationship Discovery", "Predictive Pattern Discovery"),
    ("Team Synergy", "Dataset Quality"),
    ("Market Disruption", "Revenue Trend"),
    ("Technology Migration", "Metric Change"),
    ("Microsoft partnered with OpenAI during Build 2024.", "Sales data shows Product A revenue increasing from January to March."),
    ("Microsoft partnered with OpenAI.", "Product A revenue increased across the quarter."),
    ("Microsoft", "Product A"),
    ("Build 2024", "March sales"),
    ("Who collaborated with OpenAI?", "What is the total revenue by product?"),
    ("Who partnered with OpenAI?", "Which product generated the highest revenue?"),
    ("Which technologies are mentioned?", "Which columns have missing values?"),
    ("What events are connected to Microsoft?", "What trends appear in the uploaded dataset?"),
    ("Project Alpha is cancelled", "a high-revenue product declines"),
    ("AWS is replaced by Azure", "missing values are cleaned"),
    ("Person A", "Dataset A"),
    ("Project Z", "Metric Z"),
    ("Team A and Team B are connected through Project X", "Dataset A and Dataset B share similar performance patterns"),
    ("but currently have no direct relationship.", "but require comparison through selected metrics."),
    ("PARTNERED_WITH", "CORRELATES_WITH"),
    ("WORKS_ON", "CONTAINS_METRIC"),
    ("USES_TECHNOLOGY", "USES_COLUMN"),
    ("PRESENTED_AT", "GENERATED_CHART"),
    ("PART_OF", "BELONGS_TO_DATASET"),
    ("COLLABORATED_WITH", "COMPARED_WITH"),
    ("MENTIONS", "REFERENCES"),
    ("Person", "User"),
    ("Organization", "Dataset"),
    ("Technology", "Metric"),
    ("Event", "Analysis"),
    ("People", "Users"),
    ("Organizations", "Datasets"),
    ("Technologies", "Metrics"),
    ("Events", "Analyses"),
    ("Corporate Knowledge Management", "Business Dataset Analysis"),
    ("Dependency & Impact Analysis", "Revenue and Operations Analysis"),
    ("Cross-Team Discovery", "Cross-Dataset Insight Discovery"),
    ("/api/auth/register", "/users/"),
    ("/api/auth/login", "/token"),
    ("/api/upload", "/datasets/upload"),
    ("/api/extract", "/datasets/{dataset_id}/profile"),
    ("/api/documents/delete", "/datasets/{dataset_id}"),
    ("/api/documents", "/datasets/"),
    ("/api/graph/relationships", "/chat/history/{session_id}"),
    ("/api/graph/nodes", "/datasets/{dataset_id}/data"),
    ("/api/graph", "/datasets/{dataset_id}/profile"),
    ("/api/stats", "/admin/stats"),
    ("/api/chat", "/chat/"),
    ("/api/discover", "/datasets/{dataset_id}/sentiment"),
    ("/api/simulation/types", "/admin/logs"),
    ("/api/simulation/run", "/datasets/{dataset_id}/clean"),
    ("/api/simulation/predict-future", "/datasets/{dataset_id}/summary"),
    ("/api/simulation/report", "/chat/"),
    ("/api/health", "/health"),
    ("All API routes are prefixed with /api.", "FastAPI exposes REST routes directly from the application root."),
    ("All API endpoints are prefixed with /api.", "The API is available from the configured backend base URL."),
    ("cd omnigraph-backend", "cd backend"),
    ("cd omnigraph-frontend", "cd frontend"),
    ("npm install", "pip install -r requirements.txt"),
    ("npm run dev", "uvicorn main:app --reload"),
    ("PORT=5000", "DATABASE_URL=sqlite:///./sql_app.db"),
    ("DATABASE_URL=postgresql://user:pass@host/omnigraph", "SECRET_KEY=your_secret_key"),
    ("NEO4J_URI=bolt://localhost:7687", "CHROMA_HOST=localhost"),
    ("NEO4J_USER=neo4j", "CHROMA_TENANT=default_tenant"),
    ("NEO4J_PASSWORD=your_neo4j_password", "CHROMA_DATABASE=default_database"),
    ("CHROMA_HOST=bolt://localhost:7687", "CHROMA_HOST=localhost"),
    ("CHROMA_TENANT=neo4j", "CHROMA_TENANT=default_tenant"),
    ("CHROMA_API_KEY=your_neo4j_password", "CHROMA_DATABASE=default_database"),
    ("OPENAI_API_KEY=sk-...", "GOOGLE_API_KEY=your_google_key"),
    ("LLM_PROVIDER=openai   # or groq", "GROQ_API_KEY=your_groq_key"),
    ("npx prisma migrate dev --name init", "python -m uvicorn main:app --reload"),
    ("Server runs on http://localhost:5000", "Backend runs on http://localhost:8000"),
    ("VITE_API_URL=http://localhost:5000", "NEXT_PUBLIC_API_URL=http://localhost:8000"),
    ("Frontend runs on http://localhost:5173", "Frontend runs on http://localhost:3000"),
    ("MATCH (n) RETURN n LIMIT 25", "GET http://localhost:8000/health"),
    ("Executive Knowledge Platform", "Executive Analytics Platform"),
]


def transform_text(text: str) -> str:
    updated = text
    for old, new in REPLACEMENTS:
        updated = updated.replace(old, new)
    return updated


def set_paragraph_text_keep_style(paragraph, text: str) -> None:
    if paragraph.text == text:
        return
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    first = runs[0]
    for run in runs[1:]:
        run.text = ""
    first.text = text


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


def rewrite_text_preserving_layout(path: Path) -> None:
    doc = Document(path)
    doc.core_properties.title = f"{PROJECT} Project Documentation"
    doc.core_properties.subject = "Technical project documentation"
    doc.core_properties.author = STUDENT
    doc.core_properties.keywords = "AI data analyst, FastAPI, Next.js, LangGraph, ChromaDB, RAG, data analytics"

    for paragraph in doc.paragraphs:
        if paragraph.text:
            set_paragraph_text_keep_style(paragraph, transform_text(paragraph.text))

    for table in doc.tables:
        for paragraph in iter_table_paragraphs(table):
            if paragraph.text:
                set_paragraph_text_keep_style(paragraph, transform_text(paragraph.text))

    # Fix key table values that should be exact rather than phrase-replaced.
    if doc.tables:
        t = doc.tables[0]
        exact_rows = [
            ("Project Name", PROJECT),
            ("Project Type", "AI-Powered Data Analytics and Natural Language BI Platform"),
            ("Student", STUDENT),
            ("Roll Number", ROLL),
            ("Tech Stack", "FastAPI + SQLAlchemy | Next.js + TypeScript | SQLite/PostgreSQL-ready | ChromaDB | LangGraph | Gemini/Groq/Mistral"),
            ("Document Type", "Technical Project Documentation"),
        ]
        for row, (left, right) in zip(t.rows, exact_rows):
            set_paragraph_text_keep_style(row.cells[0].paragraphs[0], left)
            set_paragraph_text_keep_style(row.cells[1].paragraphs[0], right)

    doc.save(path)


def font(size: int, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\segoeuib.ttf" if bold else r"C:\Windows\Fonts\segoeui.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_wrapped(draw, xy, text, fnt, fill, width, spacing=6):
    x, y = xy
    avg = max(8, int(fnt.size * 0.55))
    chars = max(12, width // avg)
    lines = []
    for para in text.split("\n"):
        lines.extend(textwrap.wrap(para, width=chars) or [""])
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        y += fnt.size + spacing
    return y


def card(draw, xyxy, fill="#ffffff", outline="#d8e3f0", radius=18, width=2):
    draw.rounded_rectangle(xyxy, radius=radius, fill=fill, outline=outline, width=width)


def make_visual(size, title, subtitle, variant):
    w, h = size
    img = Image.new("RGB", size, "#f7fafc")
    d = ImageDraw.Draw(img)

    # Header band
    d.rectangle([0, 0, w, max(90, h // 9)], fill="#12355b")
    d.text((w * 0.04, h * 0.035), "AI Data Analyst Agent", font=font(max(18, w // 55), True), fill="#ffffff")
    d.text((w * 0.04, h * 0.075), title, font=font(max(12, w // 90)), fill="#b9d7f2")

    margin = max(18, w // 35)
    top = max(115, h // 7)
    accent = ["#2563eb", "#16a34a", "#7c3aed", "#dc2626", "#0891b2"][variant % 5]

    if variant == 0:
        # Dashboard
        cols = 4 if w > 900 else 2
        card_w = (w - 2 * margin - (cols - 1) * 18) // cols
        card_h = max(90, h // 7)
        labels = [("Total Datasets", "12"), ("Storage Used", "248 KB"), ("Avg Size", "20.7 KB"), ("AI Engine", "Online")]
        for i, (a, b) in enumerate(labels):
            x = margin + (i % cols) * (card_w + 18)
            y = top + (i // cols) * (card_h + 18)
            card(d, [x, y, x + card_w, y + card_h])
            d.text((x + 18, y + 18), a, font=font(max(12, w // 105), True), fill="#334155")
            d.text((x + 18, y + 48), b, font=font(max(22, w // 50), True), fill=accent)
        y2 = top + ((len(labels) + cols - 1) // cols) * (card_h + 24)
        card(d, [margin, y2, w - margin, h - margin], fill="#eef6ff", outline="#bad7f4")
        d.text((margin + 24, y2 + 22), "AI Data Ingestion", font=font(max(18, w // 65), True), fill="#1e3a8a")
        draw_wrapped(d, (margin + 24, y2 + 60), "Upload CSV/Excel files, launch profiling, and open recent analytics suites from one dashboard.", font(max(12, w // 95)), "#334155", w - 2 * margin - 48)
    elif variant == 1:
        # Architecture
        boxes = [
            ("Next.js Frontend", "#dbeafe"),
            ("FastAPI Backend", "#dcfce7"),
            ("LangGraph Agent", "#ede9fe"),
            ("SQLAlchemy DB", "#fef3c7"),
            ("ChromaDB RAG", "#cffafe"),
            ("LLM + Pandas", "#fee2e2"),
        ]
        bw = (w - 2 * margin - 40) // 3 if w > 900 else w - 2 * margin
        bh = max(70, h // 8)
        for i, (label, fill) in enumerate(boxes):
            x = margin + (i % 3) * (bw + 20) if w > 900 else margin
            y = top + (i // 3) * (bh + 45) if w > 900 else top + i * (bh + 20)
            card(d, [x, y, x + bw, y + bh], fill=fill)
            d.text((x + 20, y + bh // 2 - 12), label, font=font(max(14, w // 82), True), fill="#172033")
        d.text((margin, h - margin - 45), "Flow: upload -> profile -> embed context -> ask -> analyze -> chart/report", font=font(max(12, w // 95), True), fill="#334155")
    elif variant == 2:
        # Dataset table
        card(d, [margin, top, w - margin, h - margin], fill="#ffffff")
        headers = ["Column", "Type", "Missing", "Unique", "Mean"]
        rows = [
            ["revenue", "float", "0", "42", "5840.25"],
            ["department", "text", "2", "6", "-"],
            ["salary", "int", "1", "38", "61200"],
            ["feedback", "text", "0", "118", "-"],
        ]
        x0, y0 = margin + 24, top + 64
        colw = (w - 2 * margin - 48) // len(headers)
        d.text((x0, top + 24), subtitle, font=font(max(16, w // 72), True), fill="#1e3a8a")
        for i, head in enumerate(headers):
            d.rectangle([x0 + i * colw, y0, x0 + (i + 1) * colw, y0 + 42], fill="#1f4e79")
            d.text((x0 + i * colw + 10, y0 + 11), head, font=font(max(10, w // 110), True), fill="#ffffff")
        for r, row in enumerate(rows):
            for i, val in enumerate(row):
                yy = y0 + 42 + r * 48
                d.rectangle([x0 + i * colw, yy, x0 + (i + 1) * colw, yy + 48], outline="#d8e3f0", fill="#f8fbff" if r % 2 else "#ffffff")
                d.text((x0 + i * colw + 10, yy + 15), val, font=font(max(10, w // 115)), fill="#334155")
    elif variant == 3:
        # Chat mobile/tall or chat panel
        card(d, [margin, top, w - margin, h - margin], fill="#ffffff")
        y = top + 28
        bubbles = [
            ("User", "Plot revenue by product", "#dbeafe"),
            ("AI", "Product B has the highest revenue. Here is an interactive Plotly chart.", "#f1f5f9"),
            ("User", "Any missing values?", "#dbeafe"),
            ("AI", "Two department values are missing and can be filled during cleaning.", "#f1f5f9"),
        ]
        for who, msg, fill in bubbles:
            bw = w - 2 * margin - 46
            bh = max(70, int(h * 0.09))
            x = margin + 22
            card(d, [x, y, x + bw, y + bh], fill=fill, outline="#d8e3f0", radius=14)
            d.text((x + 14, y + 10), who, font=font(max(10, w // 120), True), fill=accent)
            draw_wrapped(d, (x + 14, y + 30), msg, font(max(10, w // 130)), "#334155", bw - 28, spacing=3)
            y += bh + 18
    else:
        # Chart/admin visual
        card(d, [margin, top, w - margin, h - margin], fill="#ffffff")
        if h < 260 or w < 520:
            d.ellipse([margin + 18, top + 22, margin + 58, top + 62], fill="#16a34a")
            d.text((margin + 76, top + 18), "AI Engine", font=font(max(14, w // 20), True), fill="#1e3a8a")
            d.text((margin + 76, top + 54), "Online", font=font(max(12, w // 26), True), fill="#16a34a")
            return img
        chart_x, chart_y = margin + 60, top + 80
        chart_w, chart_h = w - 2 * margin - 120, h - top - margin - 150
        d.line([chart_x, chart_y + chart_h, chart_x + chart_w, chart_y + chart_h], fill="#94a3b8", width=3)
        d.line([chart_x, chart_y, chart_x, chart_y + chart_h], fill="#94a3b8", width=3)
        vals = [0.35, 0.55, 0.46, 0.82, 0.68]
        bw = chart_w // (len(vals) * 2)
        for i, v in enumerate(vals):
            x = chart_x + 40 + i * bw * 2
            y = chart_y + chart_h - int(chart_h * v)
            d.rounded_rectangle([x, y, x + bw, chart_y + chart_h], radius=8, fill=accent)
        d.text((margin + 42, top + 30), subtitle, font=font(max(16, w // 72), True), fill="#1e3a8a")
        d.text((margin + 42, h - margin - 58), "Generated from selected datasets with explainable AI commentary.", font=font(max(12, w // 95)), fill="#334155")

    return img


IMAGE_SPECS = {
    "word/media/image1.png": ("Dashboard Overview", "Upload, metrics, and recent activity", 0),
    "word/media/image2.png": ("System Architecture", "Client, API, intelligence, and data layers", 1),
    "word/media/image3.png": ("Dataset Library", "Owned datasets, size, status, and actions", 2),
    "word/media/image4.png": ("Analysis Chat", "Mobile-style conversational analytics", 3),
    "word/media/image5.png": ("AI Engine Status", "Online analytics core", 4),
    "word/media/image6.png": ("Dataset Preview", "Tall raw-row inspection view", 2),
    "word/media/image7.png": ("Login and Protected Dashboard", "JWT-backed user access", 0),
    "word/media/image8.png": ("Dataset Analytics Suite", "Profile table, summary, and export", 2),
    "word/media/image9.png": ("Interactive Plotly Chart", "AI-generated visualization output", 4),
    "word/media/image10.png": ("Admin Audit Dashboard", "Users, datasets, queries, and logs", 4),
    "word/media/image11.png": ("Deployment and API Docs", "Docker, FastAPI docs, and local run flow", 1),
}


def replacement_media_bytes(original_docx: Path):
    result = {}
    with ZipFile(original_docx) as zf:
        for name, (title, subtitle, variant) in IMAGE_SPECS.items():
            with Image.open(BytesIO(zf.read(name))) as im:
                generated = make_visual(im.size, title, subtitle, variant)
            out = BytesIO()
            generated.save(out, format="PNG", optimize=True)
            result[name] = out.getvalue()
    return result


def replace_media(path: Path) -> None:
    media = replacement_media_bytes(SOURCE)
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
    with ZipFile(path, "r") as zin, ZipFile(tmp_path, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in media:
                data = media[item.filename]
            zout.writestr(item, data)
    shutil.move(str(tmp_path), path)


def main():
    shutil.copy2(SOURCE, OUTPUT)
    rewrite_text_preserving_layout(OUTPUT)
    replace_media(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
